import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(title, student_name, university, company, specialty, supervisor, academic_year, logo_univ, logo_comp, sections_dict, section_options, selected_ids, output_path):
    doc = Document()
    
    # Logos Table: School on Left, Company on Right
    table = doc.add_table(rows=1, cols=3)
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(3.0)
    table.columns[2].width = Inches(1.5)
    
    cells = table.rows[0].cells
    
    if logo_univ and os.path.exists(logo_univ):
        try:
            p = cells[0].paragraphs[0]
            r = p.add_run()
            r.add_picture(logo_univ, width=Inches(1.2))
        except: pass
        
    if logo_comp and os.path.exists(logo_comp):
        try:
            p = cells[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run()
            r.add_picture(logo_comp, width=Inches(1.2))
        except: pass

    doc.add_paragraph() # Spacer after logos
    
    if university:
        p = doc.add_paragraph()
        run = p.add_run(university)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph() # Spacer

    p = doc.add_paragraph()
    run = p.add_run("RAPPORT DE PROJET DE FIN D'ÉTUDES")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer

    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer
    
    # Info list
    if student_name:
        p = doc.add_paragraph()
        p.add_run("Réalisé par : ").bold = True
        p.add_run(student_name)
        
    if specialty:
        p = doc.add_paragraph()
        p.add_run("Spécialité / Filière : ").bold = True
        p.add_run(specialty)
        
    if company:
        p = doc.add_paragraph()
        p.add_run("Entreprise d'accueil : ").bold = True
        p.add_run(company)
        
    if supervisor:
        p = doc.add_paragraph()
        p.add_run("Encadré par : ").bold = True
        p.add_run(supervisor)

    doc.add_paragraph()
    doc.add_paragraph()
    
    if academic_year:
        p = doc.add_paragraph()
        run = p.add_run(f"Année universitaire : {academic_year}")
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_page_break()

    # Sections
    for idx, sid in enumerate(selected_ids):
        if sid in sections_dict:
            section_title = section_options[sid]['title']
            doc.add_heading(section_title, level=1)
            
            content = sections_dict[sid]
            paragraphs = content.split('\n')
            
            for line in paragraphs:
                line = line.strip()
                if not line:
                    continue
                
                # Headings
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('- ') or line.startswith('* '):
                    # List bullet
                    p = doc.add_paragraph(style='List Bullet')
                    line_text = line[2:].strip()
                    _add_markdown_runs(p, line_text)
                else:
                    # Normal paragraph
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    _add_markdown_runs(p, line)
                    
            if idx < len(selected_ids) - 1:
                # Page break for all sections except the last one
                # Usually we want a page break for new sections like chapters
                if sid in ["remerciements", "intro_generale", "chapitre_1", "chapitre_2", "chapitre_3", "conclusion"]:
                    doc.add_page_break()
                    
    doc.save(output_path)

def _add_markdown_runs(paragraph, text):
    """
    Simplement parser **bold** and *italic* pour python-docx.
    """
    # Parse bold
    parts_bold = re.split(r'(\*\*.*?\*\*)', text)
    for part_b in parts_bold:
        if part_b.startswith('**') and part_b.endswith('**'):
            # It's bold. Check if it contains italic inside
            core_b = part_b[2:-2]
            parts_both = re.split(r'(\*.*?\*)', core_b)
            for part_both in parts_both:
                if part_both.startswith('*') and part_both.endswith('*'):
                    run = paragraph.add_run(part_both[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(part_both)
                    run.bold = True
        else:
            # Not bold. Check for italic
            parts_italic = re.split(r'(\*.*?\*)', part_b)
            for part_i in parts_italic:
                if part_i.startswith('*') and part_i.endswith('*'):
                    run = paragraph.add_run(part_i[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(part_i)
